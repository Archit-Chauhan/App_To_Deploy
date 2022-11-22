import json
import streamlit as st
from streamlit_option_menu import option_menu
from PIL import Image
from docx import Document
from docx.shared import Inches
from PIL import Image
from streamlit_lottie import st_lottie

with Image.open("favicon2.png") as icon:
    st.set_page_config(page_title="Get_Resume",page_icon=icon,layout="wide")


    
def load_lottiefile(filepath: str):
    with open(filepath, "r") as f:
        return json.load(f)
lottie_man = load_lottiefile("man.json")
lottie_contact = load_lottiefile("contact.json")
lottie_done = load_lottiefile('done.json')



col1,col2=st.columns(2)
with col1:
    st.title('# **_RESUME BUILDER_**')
with col2:
    st_lottie(
        lottie_man,
        height=200,
        width=200
    )
# col1,col2=st.columns(2)
# with col1:
#     st_lottie(
#         lottie_coding,
#         height=150,
#         width=150
#     )
# with col2:
#     st_lottie(
#         lottie_coding2,
#         height=150,
#         width=150
#     )


# <--------------------------------entry checkers-------------------------------->
def checkname(fname):
    if(fname==''):
        st.error("Enter Your Name")
        return 0
    else:
        return 1


def checkiarea(iarea):
    if(iarea==''):
        st.error("Enter Your Major Interest")
        return 0
    else:
        return 1


def checkiareaexp(iareaexp):
    if(iareaexp==0):
        st.error("Select Your Experience In Major Interest Area")
        return 0
    else:
        return 1


def checkcollege(college):
    if(college==""):
        st.error("Enter Your College")
        return 0
    else:
        return 1


def checkgradyear(gradyear):
    if(gradyear=='<select>'):
        print("Select Your Graduation Year")
        return 0
    else:
        return 1

def checkskills(skills):
    if(len(skills)==0):
        print("Select Your Skills")
        return 0
    else:
        return 1
#<-------------------------------------------------checkers end-------------------------------------->



#<--------------------------------------------------submit-------------------------------------------->

def submit(fname,sname,defineu,iarea,iareaexp,college,gradyear,skills):
    file = Document()
    file.add_heading((" ".join([fname.capitalize(), sname.capitalize()]))  , 0)
    file.add_paragraph("Student of {} currrently in {} year of B.Tech and a keen learner.{}.My learnings are as follows:".format(college,gradyear,defineu))
    para=file.add_heading("My Interested Area is {}".format(iarea),2)
    para.paragraph_format.line_spacing = Inches(0.5)
    file.add_paragraph("Having experience in the same of {} years".format(iareaexp))
    file.add_paragraph("___")
    file.add_paragraph('Skills', style='Intense Quote')
    for i in skills:
        file.add_paragraph(
    i, style='List Bullet'
)
    file.save("generatedLetter.docx")
    with open("generatedLetter.docx", "rb") as file:
        col1,col2=st.columns(2)
        with col1:
            btn = st.download_button(
                label="Download Now",
                data=file,
                file_name=fname+"resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
              )
        
        with col2:
            st_lottie(
                lottie_done,
                width=300,
                height=300
            )
        if(btn):
            st.balloons
#<-----------------------------------------submit ends -------------------------------------->

def main():                                                               #main page
    fname = st.text_input('First name*')
    checkreturn1=checkname(fname)


    sname = st.text_input('Last name')

    defineu=st.text_input("Paragragh That Defines You Well")

    col1, col2 = st.columns(2)
    iarea = col1.text_input('Your Major Interest Area*')
    iareaexp = col2.slider('Select Your Experience in the same', 0, 10)
    checkreturn2=checkiarea(iarea)
    checkreturn3=checkiareaexp(iareaexp)


    college=st.text_area("Your College Name*")
    checkreturn4=checkcollege(college)


    gradyear = st.selectbox(
        'Select Year*', ['First', 'Second', 'Third', 'Fourth'],3)
    checkreturn4=checkgradyear(gradyear)
    skills = st.multiselect("Select Skills*:", [
                            'Programming', 'Data science', 'Digital Marketing', 'Cyber Security', 'Project Management'])
    checkreturn5=checkskills(skills)
    if(checkreturn1 and checkreturn2 and checkreturn3 and checkreturn4):
        submit(fname,sname,defineu,iarea,iareaexp,college,gradyear,skills)








def menu():
    menu_select = option_menu(
        menu_title="Main Menu",
        options=["Home", "About", "Contact Us"],
        icons=["house", "book", "contact"],
        menu_icon="cast",
        orientation="horizontal"
    )
    if menu_select == "Home":
        main()
    if menu_select =="About":
        lottie_team = load_lottiefile("team.json")
        col1,col2=st.columns(2)
        with col1:
            st_lottie(
                lottie_team,
                height=300,
                width=300
            )
        with col2:
            st.markdown("""
            
            ## Team Members :-
            1. Archit Chauhan (2101430100036)
            2. Raghav Jindal (2101430100138)
            3. Sakshi Priya (2101430100153)
            4. Vanshika Diwedi (2101430100190)
            5. Lovely Agrawal (2101430100097)
            """)
    if menu_select == "Contact Us":
        col1,col2 = st.columns(2)
        with col1:
            st_lottie(
                lottie_contact,
                height=300,
                width=300
            )
        with col2:
            st.markdown("""
            ## YOU CAN CONTACT ME -                 
            Email Id : architchauhan.12b1.10@gmail.com
            """)



def hideFooter():
    hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {
	            visibility: hidden;
            }
            footer:after {
                content:'Made By Developers'; 
                visibility: visible;
                display: block;
                position: relative;
                #background-color: red;
                padding: 5px;
                top: 2px;
            }
            }
                        </style>
                        """
    st.markdown(hide_streamlit_style, unsafe_allow_html=True)

hideFooter()

menu()

